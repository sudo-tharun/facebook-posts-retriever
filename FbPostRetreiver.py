#importing packages
import json
from translate import Translator
from datetime import datetime
import docx
doc=docx.Document()
da = json.load(open("Enter the name of the JSON file here: ")) #loading JSON file
i=0
s=datetime(2019,12,22,00,00,00) #sort according to date
for i in range(len(da)):
    try:
        timestamp = da[i]['timestamp']
        dt = datetime.fromtimestamp(timestamp)
        if dt>s: #comparing dates
            k=da[i]['data'][0]['post'] #Specifying the key values in the dictionaries.
            k = k.encode('latin-1').decode('unicode_escape').encode('latin-1').decode('utf-8') #decoding the levels of encryption
            doc.add_heading(str(dt)) #writing date and time as the heading to the word document
            doc.add_paragraph(k) #adding post's text as paragraph in the word document
    except:
        pass
doc.save("FBPostsRetrieve.docx") #File name is arbitrary. Saving File...

#This is only applicable when you wish to write only POSTS from your timeline. Follow the instructions in the README.md file.
